"""Turn an uploaded resume into embeddable chunks.

Section-aware rather than fixed-size. A resume is not prose: splitting every
600 characters routinely cuts a bullet in half, so the fragment that should have
evidenced "5 years of Kafka" becomes two fragments that evidence neither. Chunks
follow the document's own structure instead, and stay whole.
"""

import io
import re
from dataclasses import dataclass
from datetime import date
from typing import Any

from app.core.exceptions import InvalidOperationError
from app.domain.enums import Seniority

MAX_UPLOAD_BYTES = 5 * 1024 * 1024
MIN_TEXT_LENGTH = 200

# Headings a resume actually uses, mapped to a canonical section. The label is
# carried onto each chunk so retrieval can weight experience above education
# without re-reading the text.
SECTION_PATTERNS: list[tuple[str, re.Pattern[str]]] = [
    (
        "experience",
        re.compile(
            r"^\s*(work|professional|relevant|industry)?\s*"
            r"(experience|employment|employment\s+history|career\s+history|"
            r"work\s+history|internships?)\b",
            re.I,
        ),
    ),
    ("projects", re.compile(r"^\s*(projects?|personal\s+projects?|selected\s+work)\b", re.I)),
    (
        "skills",
        re.compile(r"^\s*(technical\s+)?(skills|technologies|tech\s+stack|competencies)\b", re.I),
    ),
    ("education", re.compile(r"^\s*(education|academics?|qualifications?)\b", re.I)),
    ("certifications", re.compile(r"^\s*(certifications?|licenses?|courses?)\b", re.I)),
    ("summary", re.compile(r"^\s*(summary|profile|objective|about\s+me)\b", re.I)),
    ("publications", re.compile(r"^\s*(publications?|papers?|research)\b", re.I)),
    ("awards", re.compile(r"^\s*(awards?|achievements?|honou?rs?)\b", re.I)),
]

# A heading is short and title-like. Length alone is unreliable, so it must also
# match a known pattern before a line is treated as one.
MAX_HEADING_LENGTH = 60

# Bullets shorter than this ("Python", "•") carry no retrievable meaning on
# their own and are merged into the previous chunk rather than embedded.
MIN_CHUNK_LENGTH = 40
MAX_CHUNK_LENGTH = 1200


@dataclass(frozen=True, slots=True)
class Chunk:
    ordinal: int
    section: str | None
    content: str


def extract_text(filename: str, data: bytes) -> str:
    """Pull plain text out of a PDF or DOCX upload."""
    if len(data) > MAX_UPLOAD_BYTES:
        raise InvalidOperationError(
            f"That file is {len(data) / 1_048_576:.1f} MB; the limit is "
            f"{MAX_UPLOAD_BYTES // 1_048_576} MB. A resume should be far smaller — "
            "check it is not a scanned image."
        )

    lowered = filename.lower()
    if lowered.endswith(".pdf"):
        text = _extract_pdf(data)
    elif lowered.endswith((".docx", ".doc")):
        text = _extract_docx(data)
    elif lowered.endswith((".txt", ".md")):
        text = data.decode("utf-8", errors="replace")
    else:
        raise InvalidOperationError("Upload a PDF, DOCX, or plain text file.")

    text = _normalize(text)

    if len(text) < MIN_TEXT_LENGTH:
        raise InvalidOperationError(
            "Almost no text could be read from that file. If it is a scanned "
            "PDF the words are an image — export a text-based PDF from your "
            "editor, or paste the text instead."
        )
    return text


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    try:
        reader = PdfReader(io.BytesIO(data))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception as exc:  # pypdf raises a wide variety on malformed files
        raise InvalidOperationError(f"That PDF could not be read: {exc}") from exc


def _extract_docx(data: bytes) -> str:
    import docx

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:
        raise InvalidOperationError(f"That document could not be read: {exc}") from exc

    parts = [p.text for p in document.paragraphs]
    # Resumes frequently lay out experience in invisible tables, and those cells
    # hold the actual content. Skipping them loses most of the document.
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def _normalize(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\xa0", " ")
    text = re.sub(r"[ \t]+", " ", text)
    # Normalise the several bullet glyphs PDFs emit to one marker, so the
    # bullet-detection below does not depend on which tool produced the file.
    text = re.sub(r"^[\s]*[•▪◦‣∙·–—*]\s*", "- ", text, flags=re.M)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def detect_section(line: str) -> str | None:
    stripped = line.strip()
    if not stripped or len(stripped) > MAX_HEADING_LENGTH:
        return None
    for name, pattern in SECTION_PATTERNS:
        if pattern.match(stripped):
            return name
    return None


def chunk_resume(text: str) -> list[Chunk]:
    """Split into retrievable passages, tagged by section.

    Each bullet or paragraph becomes its own chunk, because the unit that
    answers "does this candidate have Kafka experience?" is a single
    accomplishment line, not a page. Very short lines merge into the previous
    chunk, and over-long paragraphs are split on sentence boundaries.
    """
    chunks: list[Chunk] = []
    section: str | None = None
    buffer: list[str] = []

    def flush() -> None:
        if not buffer:
            return
        content = " ".join(part.strip() for part in buffer if part.strip()).strip()
        buffer.clear()
        if not content:
            return

        for piece in _split_long(content):
            if (
                chunks
                and len(piece) < MIN_CHUNK_LENGTH
                and chunks[-1].section == section
                and len(chunks[-1].content) + len(piece) < MAX_CHUNK_LENGTH
            ):
                merged = f"{chunks[-1].content} {piece}"
                chunks[-1] = Chunk(chunks[-1].ordinal, section, merged)
            else:
                chunks.append(Chunk(len(chunks), section, piece))

    for line in text.split("\n"):
        heading = detect_section(line)
        if heading:
            flush()
            section = heading
            continue

        if not line.strip():
            flush()
            continue

        # A new bullet ends the previous one; a continuation line extends it.
        if line.lstrip().startswith("- ") and buffer:
            flush()

        buffer.append(line)

    flush()
    return chunks


def _split_long(content: str) -> list[str]:
    if len(content) <= MAX_CHUNK_LENGTH:
        return [content]

    sentences = re.split(r"(?<=[.!?])\s+", content)
    pieces: list[str] = []
    current = ""
    for sentence in sentences:
        if len(current) + len(sentence) + 1 > MAX_CHUNK_LENGTH and current:
            pieces.append(current.strip())
            current = sentence
        else:
            current = f"{current} {sentence}".strip()
    if current:
        pieces.append(current.strip())
    return pieces


YEARS_PATTERNS = [
    re.compile(r"(\d+(?:\.\d+)?)\s*\+?\s*years?\s+of\s+(?:professional\s+)?experience", re.I),
    re.compile(r"(?:experience|exp)\s*[:\-]\s*(\d+(?:\.\d+)?)\s*\+?\s*years?", re.I),
]

MAX_PLAUSIBLE_YEARS = 50


def stated_years_experience(text: str) -> float | None:
    """Years of experience as the resume itself claims them, if it does."""
    for pattern in YEARS_PATTERNS:
        match = pattern.search(text)
        if match:
            value = float(match.group(1))
            if 0 < value <= MAX_PLAUSIBLE_YEARS:
                return value
    return None


# --- Employment history ----------------------------------------------------
#
# Most resumes never write "4 years of experience" anywhere; they write the
# dates and leave the arithmetic to the reader. Leaving it unknown is not
# neutral — `score_experience` and `score_seniority` both return 0.5 for an
# unknown candidate, so a quarter of every match score becomes a constant.
#
# The three ways inferring from dates goes wrong are each handled rather than
# assumed away: concurrent roles are merged so they are counted once, education
# and personal-project dates are excluded, and only role header lines are read
# so that "cut latency from 2020 to 2022" in a bullet cannot become a job. An
# explicit statement still wins where there is one, and the estimate records
# which of the two it came from.

_MONTHS = {
    "jan": 1, "feb": 2, "mar": 3, "apr": 4, "may": 5, "jun": 6,
    "jul": 7, "aug": 8, "sep": 9, "oct": 10, "nov": 11, "dec": 12,
}  # fmt: skip

_MONTH_RE = r"(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?"
_YEAR_RE = r"(?:19|20)\d{2}"
# "March 2019", "Mar. 2019", "03/2019", "2019"
_DATE_RE = rf"(?:{_MONTH_RE}[\s.,'’-]*)?(?:\d{{1,2}}\s*[/.]\s*)?{_YEAR_RE}"
_PRESENT_RE = r"present|current(?:ly)?|now|today|ongoing|date"
_SEPARATOR_RE = r"\s*(?:[–—‒-]|\bto\b|\buntil\b|\bthrough\b)\s*"

DATE_RANGE = re.compile(
    rf"(?<![\w/])(?P<start>{_DATE_RE}){_SEPARATOR_RE}(?P<end>{_DATE_RE}|{_PRESENT_RE})(?!\w)",
    re.I,
)

# Degrees and institutions. Dotted forms only where the dotless one is an
# ordinary English word — "\bb\.?e\.?\b" would match every "be" in the document.
_EDUCATION_HINTS = re.compile(
    r"\b(b\.tech|btech|b\.e\.|b\.sc|bsc|bachelors?|m\.tech|mtech|m\.sc|msc|masters?|"
    r"mba|ph\.?d|doctorate|diploma|university|college|institute|academy|"
    r"high\s+school|g\.?p\.?a|c\.?g\.?p\.?a)\b",
    re.I,
)

_ROLE_WORDS = re.compile(
    r"\b(engineer|engineering|developer|programmer|architect|scientist|analyst|"
    r"consultant|designer|manager|director|lead|head|intern|trainee|specialist|"
    r"administrator|researcher|founder|president|officer|associate|assistant|"
    r"devops|sre|qa)\b",
    re.I,
)

_COMPANY_HINTS = re.compile(
    r"\b(inc|llc|ltd|limited|gmbh|corp|corporation|company|technologies|technology|"
    r"labs?|systems?|solutions?|software|consulting|group|partners|studios?|pvt|"
    r"private|bank|media|digital|ventures)\b\.?",
    re.I,
)

# Separators a header line uses between title, company, and location. Hyphens
# are deliberately absent: they would split "Full-Stack Engineer" in two.
_HEADER_SPLIT = re.compile(r"\s*(?:[|,;•·]|—|–|\bat\b|\bfor\b|@)\s*", re.I)
_BRACKETS = re.compile(r"[()\[\]{}]")

MAX_HEADER_PART_LENGTH = 60

# Lines above the dates that may still belong to the same header. Two covers
# the title-then-company-then-dates layout a PDF produces from a right-aligned
# date; more would start reaching into the previous role's bullets.
HEADER_LOOKBACK = 2

# Seniority a title states outright, most specific first — "Senior Engineering
# Manager" is a lead, not a senior, so `lead` is tested before `senior`.
TITLE_SENIORITY: list[tuple[Seniority, re.Pattern[str]]] = [
    (Seniority.INTERN, re.compile(r"\b(intern|internship|trainee|apprentice)\b", re.I)),
    (Seniority.PRINCIPAL, re.compile(r"\b(principal|distinguished|fellow)\b", re.I)),
    (Seniority.STAFF, re.compile(r"\bstaff\b", re.I)),
    (Seniority.LEAD, re.compile(r"\b(lead|manager|director|head|vp|chief|cto)\b", re.I)),
    (Seniority.SENIOR, re.compile(r"\b(senior|sr\.?)\b", re.I)),
    (Seniority.JUNIOR, re.compile(r"\b(junior|jr\.?|graduate|entry[\s-]level)\b", re.I)),
]


def seniority_from_title(title: str) -> Seniority | None:
    """The level a job title states. ``None`` when it states none."""
    for level, pattern in TITLE_SENIORITY:
        if pattern.search(title):
            return level
    return None


@dataclass(frozen=True, slots=True)
class Position:
    """One role read out of the experience section.

    ``start`` and ``end`` are month indices (``year * 12 + month``) rather than
    dates: every question asked of them is arithmetic on whole months, and an
    index makes merging overlapping roles a comparison of two integers.
    """

    title: str | None
    company: str | None
    start: int
    end: int
    is_current: bool
    raw: str

    @property
    def months(self) -> int:
        """Inclusive of both ends — a role held only in June 2020 is one month."""
        return max(0, self.end - self.start) + 1

    def as_dict(self) -> dict[str, Any]:
        return {
            "title": self.title,
            "company": self.company,
            "start": format_month(self.start),
            "end": None if self.is_current else format_month(self.end),
            "months": self.months,
        }


def format_month(index: int) -> str:
    year = (index - 1) // 12
    return f"{year:04d}-{index - year * 12:02d}"


def _parse_endpoint(token: str, today: date) -> int | None:
    """One end of a date range, as a month index."""
    token = token.strip()
    if re.fullmatch(_PRESENT_RE, token, re.I):
        return today.year * 12 + today.month

    year_match = re.search(_YEAR_RE, token)
    if not year_match:
        return None
    year = int(year_match.group())

    month = None
    if name := re.match(_MONTH_RE, token, re.I):
        month = _MONTHS.get(name.group()[:3].lower())
    elif (numeric := re.match(r"(\d{1,2})\s*[/.]", token)) and 1 <= int(numeric.group(1)) <= 12:
        month = int(numeric.group(1))

    # A bare year is read as mid-year, not January. "2019 - 2024" is five years
    # to whoever reads the resume; anchoring both ends to January makes it six.
    return year * 12 + (month or 6)


def _clean_parts(text: str) -> list[str]:
    """A header line split into its title / company / location candidates."""
    parts = []
    for part in _HEADER_SPLIT.split(_BRACKETS.sub(" ", text)):
        part = part.strip(" \t-–—|,;:·•")
        if 2 <= len(part) <= MAX_HEADER_PART_LENGTH and re.search(r"[A-Za-z]", part):
            parts.append(part)
    return parts


def _read_title_and_company(remainder: str, preceding: list[str]) -> tuple[str | None, str | None]:
    """Best-effort title and company from a role header.

    The two orderings are about equally common ("Engineer, Acme" and "Acme —
    Engineer"), so position on the line says nothing. What does say something is
    the words: one part names a role, the other usually does not.

    ``preceding`` is the couple of lines above, nearest last. A PDF lays a
    right-aligned date out on its own line often enough that the header is
    routinely three lines — title, company, dates — with nothing but the dates
    on the line that identifies it as a role at all.
    """
    parts = _clean_parts(remainder)
    for line in reversed(preceding):
        parts.extend(_clean_parts(line))

    title = next((p for p in parts if _ROLE_WORDS.search(p)), None)

    rest = [p for p in parts if p is not title]
    company = next((p for p in rest if _COMPANY_HINTS.search(p)), None)
    if company is None:
        company = rest[0] if rest else None

    return title, company


def _read_position(line: str, preceding: list[str], today: date) -> Position | None:
    match = DATE_RANGE.search(line)
    if match is None:
        return None

    start = _parse_endpoint(match.group("start"), today)
    end_token = match.group("end").strip()
    is_current = bool(re.fullmatch(_PRESENT_RE, end_token, re.I))
    end = _parse_endpoint(end_token, today)

    if start is None or end is None or end < start:
        return None
    # A role starting after next year, or half a century ago, is a
    # misidentified number rather than a job.
    horizon = (today.year + 1) * 12
    if not (horizon - MAX_PLAUSIBLE_YEARS * 12) <= start <= horizon:
        return None

    title, company = _read_title_and_company(
        f"{line[: match.start()]} {line[match.end() :]}", preceding
    )
    return Position(
        title, company, start, min(end, today.year * 12 + today.month), is_current, line
    )


def parse_positions(text: str, *, today: date | None = None) -> list[Position]:
    """Every role in the experience section, with its dates.

    Read only from role header lines. Bullets are skipped on purpose: dates
    inside them describe the work ("cut deploy time between 2020 and 2022"),
    not a period of employment, and counting those is how an inferred total
    runs away from the truth.
    """
    today = today or date.today()
    lines = text.split("\n")

    # Where the resume marks its experience section, nothing outside it counts.
    # Where it marks none, unsectioned lines are read instead — otherwise a
    # resume without headings yields nothing at all.
    has_experience_section = any(detect_section(line) == "experience" for line in lines)

    positions: list[Position] = []
    section: str | None = None
    preceding: list[str] = []

    for line in lines:
        if heading := detect_section(line):
            section = heading
            preceding.clear()
            continue

        stripped = line.strip()
        if not stripped:
            continue

        is_bullet = stripped.startswith("- ")
        in_scope = section == "experience" if has_experience_section else section is None
        if (
            in_scope
            and not is_bullet
            # Degree dates are not employment, wherever they appear.
            and not _EDUCATION_HINTS.search(stripped)
            and (position := _read_position(stripped, preceding, today))
        ):
            positions.append(position)
            # The header is consumed; its lines must not also be read as the
            # title of the next role down.
            preceding.clear()
            continue

        # Bullets are never part of a header, and letting one in would offer
        # "led the migration" as a job title.
        if not is_bullet:
            preceding.append(stripped)
            del preceding[:-HEADER_LOOKBACK]

    return positions


def years_from_positions(positions: list[Position]) -> float | None:
    """Total time employed, counting concurrent roles once.

    A promotion listed as two entries, or a contract held alongside a staff
    job, would otherwise be added together and hand back a number well beyond
    the candidate's actual career.
    """
    spans = sorted((p.start, p.end) for p in positions)
    if not spans:
        return None

    months = 0
    start, end = spans[0]
    for next_start, next_end in spans[1:]:
        if next_start <= end + 1:  # overlapping, or adjacent with no gap
            end = max(end, next_end)
        else:
            months += end - start + 1
            start, end = next_start, next_end
    months += end - start + 1

    years = round(months / 12, 1)
    return years if 0 < years <= MAX_PLAUSIBLE_YEARS else None


@dataclass(frozen=True, slots=True)
class ExperienceEstimate:
    years: float | None
    source: str | None
    """``"stated"``, ``"dates"``, or ``None`` when neither was available.

    Carried through to the UI so a number the resume never actually claims is
    not presented as though it did.
    """


def estimate_years_experience(text: str, positions: list[Position]) -> ExperienceEstimate:
    """How much experience the candidate has, and how that was determined.

    A claim in the resume wins over the dates below it. It is what a recruiter
    reads, it is what the candidate stands behind, and it already accounts for
    the things dates cannot see — a career break, unlisted early roles.
    """
    if (stated := stated_years_experience(text)) is not None:
        return ExperienceEstimate(stated, "stated")
    if (inferred := years_from_positions(positions)) is not None:
        return ExperienceEstimate(inferred, "dates")
    return ExperienceEstimate(None, None)
